from __future__ import annotations

from datetime import datetime, timezone
from pathlib import Path

from docx import Document


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_para(doc: Document, text: str) -> None:
    doc.add_paragraph(text)


def main() -> None:
    root = Path(__file__).resolve().parents[1]
    out_path = root / "Dhruva_Codebase_Story.docx"

    doc = Document()
    add_heading(doc, "Dhruva CDM: Engineering Story and Technical Deep Dive", level=0)
    add_para(
        doc,
        f"Generated on {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M:%S UTC')}.",
    )

    add_heading(doc, "1. The Story: How the System Evolved", level=1)
    add_para(
        doc,
        "The project began as a compact orbital simulation in simulator.py, where one satellite and one debris object were propagated in a 2D Matplotlib scene. "
        "That first version already carried the core idea of conjunction assessment: estimate closest approach, decide if risk is high, and trigger an avoidance burn.",
    )
    add_para(
        doc,
        "In simulator.py, both objects were advanced with RK4 integration, the closest-approach point was predicted over a finite horizon, and if the predicted miss distance "
        "dropped below a safety threshold, a small delta-v was applied in the velocity direction. This gave a visual and intuitive proof that guidance and propagation loops were working.",
    )
    add_para(
        doc,
        "After proving the physics loop, the architecture grew into a service-oriented backend plus modern web frontend. "
        "The backend became responsible for state propagation, conjunction assessment, maneuver scheduling, blackout validation, and metrics. "
        "The frontend became responsible for map view, 3D orbit view, warning visualization, analytics, and operator workflow.",
    )
    add_para(
        doc,
        "One practical turning point was update delivery. Polling snapshots can produce stale UI moments and jitter under heavier compute load. "
        "The system moved to WebSocket streaming (/orbit), where the backend advances simulation in controlled ticks and pushes fresh snapshots continuously. "
        "That improved live behavior, especially for fast-moving warning indicators and Bullseye risk display.",
    )

    add_heading(doc, "2. Codebase Overview", level=1)
    add_para(
        doc,
        "server.py creates the FastAPI app, configures CORS, instantiates SimulationState, mounts REST routes, and serves live simulation snapshots over WebSocket.",
    )
    add_para(
        doc,
        "api/routes.py exposes APIs for telemetry ingest, TLE ingest, maneuver scheduling, step-based simulation, snapshot retrieval, seed/reset controls, time-warp controls, "
        "and CSV-driven debris loading for stress testing.",
    )
    add_para(
        doc,
        "engine/sim_state.py is the operational core. It owns the current simulation time, satellite/debris state vectors, maneuver queues, conjunction warnings, blackout tracking, "
        "graveyard logic, and objective metrics.",
    )
    add_para(
        doc,
        "engine/conjunction.py computes conjunction candidates and warning levels. It supports two filtering strategies: a spatial grid method and a custom KD-tree method.",
    )
    add_para(
        doc,
        "physics.py provides the astrodynamics backbone: two-body plus J2 perturbation dynamics, RK4 integration, state propagation, rocket-equation fuel burn, and orbital period estimation.",
    )
    add_para(
        doc,
        "The dhruva-frontend app consumes snapshot/stream data and renders mission status in 2D/3D operator-friendly visual components.",
    )

    add_heading(doc, "3. Live Update Architecture: Polling vs WebSocket Streaming", level=1)
    add_para(
        doc,
        "Polling model: client repeatedly requests snapshots at fixed intervals. While simple, this can miss fast transitions and add request overhead.",
    )
    add_para(
        doc,
        "Streaming model in this codebase: server-side loop advances simulation and pushes snapshots directly using WebSocket. "
        "In server.py, state.stream_tick(...) updates the world, then ws.send_json(state.snapshot()) publishes the latest payload.",
    )
    add_para(
        doc,
        "Operational impact: lower end-to-end latency for warning updates, smoother visual feedback, and fewer synchronization issues between compute cadence and UI refresh cadence.",
    )

    add_heading(doc, "4. Conjunction Computation Choices: Grid and KD-Tree", level=1)
    add_para(
        doc,
        "The conjunction engine uses a two-stage process: (1) candidate prefiltering in position space, then (2) predictive time-of-closest-approach evaluation.",
    )
    add_para(
        doc,
        "Grid method: debris is bucketed into 3D spatial cells. For each satellite, neighboring cells are scanned for candidates. "
        "This approach is stable and performant for streaming workloads and is selected for live frontend updates.",
    )
    add_para(
        doc,
        "KD-tree method: debris positions are indexed in a 3D KD-tree. Candidate debris for each satellite are found via radius queries. "
        "This is suitable for faster backend grading/step computations with large object counts.",
    )
    add_para(
        doc,
        "Current strategy in the code: stream_tick uses method='grid' for UI smoothness, while step() uses method='kdtree' for high-throughput backend evaluation.",
    )
    add_para(
        doc,
        "After candidate filtering, each candidate pair is propagated across a look-ahead horizon, nearest miss distance is measured, and risk is labeled SAFE/WARNING/CRITICAL by threshold.",
    )

    add_heading(doc, "5. Feature-by-Feature Functional Explanation", level=1)
    features = [
        "Telemetry ingest: POST /api/telemetry accepts direct Cartesian state vectors [x,y,z,vx,vy,vz] in km and km/s.",
        "TLE ingest: POST /api/telemetry/tle converts TLE data using SGP4 and frame transformation, then ingests resulting Cartesian state vectors.",
        "CSV debris loading: POST /api/debug/load_cdm_csv maps CDM relative RTN fields into ECI-like states anchored to in-sim satellites for large-scale stress tests.",
        "Maneuver scheduling: POST /api/maneuver/schedule validates command latency, thermal cooldown, line-of-sight constraints, and fuel feasibility before acceptance.",
        "Autonomous evasions: high-risk conjunctions can trigger auto-generated evade/recovery burns, with LOS-aware timing.",
        "Blackout awareness: LOS checks against ground stations determine blackout status and estimate recovery time.",
        "Graveyard handling: low-fuel satellites can be auto-routed into graveyard maneuvers with state/status tracking.",
        "Objective metrics: uptime percentage/score, outage satellite-seconds, total delta-v, and avoidance-per-delta-v are tracked for mission-quality scoring.",
        "Snapshot API and WebSocket stream: support both pull and push data flow for dashboards and visualization tools.",
    ]
    for f in features:
        add_para(doc, f"- {f}")

    add_heading(doc, "6. Physics.py: Every Formula and Its Meaning", level=1)
    add_para(
        doc,
        "Constants:",
    )
    add_para(doc, "- MU = 398600.4418 km^3/s^2 (Earth gravitational parameter)")
    add_para(doc, "- J2 = 1.08263e-3 (Earth oblateness coefficient)")
    add_para(doc, "- RE = 6378.137 km (Earth equatorial radius)")
    add_para(doc, "- ISP = 300 s (specific impulse assumption)")
    add_para(doc, "- G0 = 9.80665 m/s^2 (standard gravity)")

    add_para(doc, "6.1 J2 acceleration model")
    add_para(
        doc,
        "Given position vector r = [x, y, z], with magnitude |r|, the J2 perturbation acceleration is implemented as:",
    )
    add_para(
        doc,
        "scalar = (3/2) * J2 * MU * RE^2 / |r|^5",
    )
    add_para(
        doc,
        "a_x = scalar * x * (5*z^2/|r|^2 - 1)",
    )
    add_para(
        doc,
        "a_y = scalar * y * (5*z^2/|r|^2 - 1)",
    )
    add_para(
        doc,
        "a_z = scalar * z * (5*z^2/|r|^2 - 3)",
    )
    add_para(
        doc,
        "This captures the first-order effect of Earth oblateness, which causes secular and periodic orbital perturbations not present in pure two-body motion.",
    )

    add_para(doc, "6.2 Total acceleration and state derivatives")
    add_para(
        doc,
        "Central gravity acceleration: a_grav = -(MU / |r|^3) * r",
    )
    add_para(
        doc,
        "Total acceleration: a_total = a_grav + a_J2",
    )
    add_para(
        doc,
        "For state S = [x,y,z,vx,vy,vz], derivative dS/dt = [vx,vy,vz,ax,ay,az].",
    )

    add_para(doc, "6.3 RK4 numerical integration")
    add_para(
        doc,
        "k1 = f(S)",
    )
    add_para(
        doc,
        "k2 = f(S + 0.5*dt*k1)",
    )
    add_para(
        doc,
        "k3 = f(S + 0.5*dt*k2)",
    )
    add_para(
        doc,
        "k4 = f(S + dt*k3)",
    )
    add_para(
        doc,
        "S_next = S + (dt/6) * (k1 + 2*k2 + 2*k3 + k4)",
    )
    add_para(
        doc,
        "RK4 gives strong accuracy/stability for this dynamic system at practical step sizes used in simulation.",
    )

    add_para(doc, "6.4 Propagation")
    add_para(
        doc,
        "propagate(state, duration, dt) repeatedly applies rk4_step over duration/dt steps, producing future state estimates.",
    )

    add_para(doc, "6.5 Tsiolkovsky-based propellant model")
    add_para(
        doc,
        "Implemented form: m_prop = m_current * (1 - exp(-|delta_v| / (ISP * G0)))",
    )
    add_para(
        doc,
        "Here delta_v is in m/s. This computes fuel consumed for a burn and enables realistic maneuver feasibility checks.",
    )

    add_para(doc, "6.6 Orbital period estimate")
    add_para(
        doc,
        "For altitude h, semi-major axis a = RE + h, and period:",
    )
    add_para(
        doc,
        "T = 2*pi*sqrt(a^3 / MU)",
    )
    add_para(
        doc,
        "This is the standard Keplerian period relation for near-circular orbits.",
    )

    add_heading(doc, "7. Data Flow from Input to Warning", level=1)
    add_para(
        doc,
        "Step 1: Objects are ingested (direct telemetry, TLE-derived state vectors, or CSV-derived stress objects).",
    )
    add_para(
        doc,
        "Step 2: Simulation propagates all objects in time with rk4_step.",
    )
    add_para(
        doc,
        "Step 3: Candidate conjunctions are filtered by grid or KD-tree, then evaluated for TCA and miss distance.",
    )
    add_para(
        doc,
        "Step 4: Risk labels drive warning displays and optional autonomous maneuver planning.",
    )
    add_para(
        doc,
        "Step 5: Snapshot payloads and stream updates carry state, warnings, maneuvers, metrics, and blackout status to the UI.",
    )

    add_heading(doc, "8. Operational Notes for Judges and Reviewers", level=1)
    add_para(
        doc,
        "For high object-count backend testing, use KD-tree mode via step() style workflows.",
    )
    add_para(
        doc,
        "For smooth operator visualization, use stream mode where the grid prefilter protects real-time responsiveness.",
    )
    add_para(
        doc,
        "When using CDM CSV stress data, remember the relative vectors are data-driven but the absolute anchor frame in simulation is synthetic unless matched to true primary states at epoch.",
    )

    doc.save(out_path)
    print(str(out_path))


if __name__ == "__main__":
    main()
