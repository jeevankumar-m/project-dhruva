from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
section = doc.sections[0]
section.top_margin    = Inches(1.0)
section.bottom_margin = Inches(1.0)
section.left_margin   = Inches(1.15)
section.right_margin  = Inches(1.15)

# ── Colour palette ────────────────────────────────────────────────────────────
NAVY   = RGBColor(0x0A, 0x1F, 0x44)   # deep navy  – headings
BLUE   = RGBColor(0x1A, 0x6B, 0xD0)   # accent blue – subheads / bullets
GOLD   = RGBColor(0xD4, 0xA0, 0x17)   # gold        – slide titles
BLACK  = RGBColor(0x1A, 0x1A, 0x1A)   # near-black  – body text
GREY   = RGBColor(0x55, 0x55, 0x55)   # grey        – notes / captions
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)


def set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def add_hr(doc):
    p   = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pb  = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "1A6BD0")
    pb.append(bot)
    pPr.append(pb)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)


def doc_title(doc, text):
    """Cover / section title."""
    p  = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(26)
    run.font.color.rgb = NAVY


def doc_subtitle(doc, text):
    p  = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.font.size      = Pt(13)
    run.font.color.rgb = BLUE
    run.italic = True


def slide_label(doc, number: int, title: str):
    """Slide number banner + slide title."""
    # Banner paragraph
    p  = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(0)
    r  = p.add_run(f"  SLIDE {number}  ")
    r.bold            = True
    r.font.size       = Pt(10)
    r.font.color.rgb  = WHITE
    # Shade the paragraph background via XML shading on the paragraph itself
    pPr  = p._p.get_or_add_pPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  "0A1F44")
    pPr.append(shd)

    # Title
    pt = doc.add_paragraph()
    pt.paragraph_format.space_before = Pt(4)
    pt.paragraph_format.space_after  = Pt(6)
    rt = pt.add_run(title)
    rt.bold           = True
    rt.font.size      = Pt(16)
    rt.font.color.rgb = GOLD
    add_hr(doc)


def section_heading(doc, text):
    p  = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(3)
    r  = p.add_run(text)
    r.bold           = True
    r.font.size      = Pt(12)
    r.font.color.rgb = BLUE


def body(doc, text, bold_parts=None):
    """Plain body paragraph. Pass bold_parts list of substrings to bold inline."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.left_indent  = Inches(0)
    if not bold_parts:
        r = p.add_run(text)
        r.font.size      = Pt(11)
        r.font.color.rgb = BLACK
    else:
        # naive inline bold: split on bold_parts
        remaining = text
        for bp in bold_parts:
            idx = remaining.find(bp)
            if idx == -1:
                continue
            before = remaining[:idx]
            if before:
                rb = p.add_run(before)
                rb.font.size = Pt(11); rb.font.color.rgb = BLACK
            rb2 = p.add_run(bp)
            rb2.bold = True; rb2.font.size = Pt(11); rb2.font.color.rgb = NAVY
            remaining = remaining[idx + len(bp):]
        if remaining:
            rb3 = p.add_run(remaining)
            rb3.font.size = Pt(11); rb3.font.color.rgb = BLACK
    return p


def bullet(doc, text, level=0, bold_lead=None):
    """Bullet point with optional bold lead phrase."""
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent = Inches(0.3 + 0.2 * level)

    if bold_lead and text.startswith(bold_lead):
        rb = p.add_run(bold_lead)
        rb.bold = True; rb.font.size = Pt(11); rb.font.color.rgb = NAVY
        rest = p.add_run(text[len(bold_lead):])
        rest.font.size = Pt(11); rest.font.color.rgb = BLACK
    else:
        r = p.add_run(text)
        r.font.size      = Pt(11)
        r.font.color.rgb = BLACK


def code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Inches(0.3)
    r = p.add_run(text)
    r.font.name       = "Courier New"
    r.font.size       = Pt(9)
    r.font.color.rgb  = RGBColor(0x10, 0x60, 0x20)
    # light grey shading
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  "F0F4F8")
    pPr.append(shd)


def punchline(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(8)
    p.paragraph_format.left_indent  = Inches(0.3)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  "EBF3FB")
    pPr.append(shd)
    r = p.add_run(text)
    r.bold           = True
    r.italic         = True
    r.font.size      = Pt(11)
    r.font.color.rgb = NAVY


def note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(text)
    r.italic         = True
    r.font.size      = Pt(10)
    r.font.color.rgb = GREY


def add_table(doc, headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    # header row
    hrow = t.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        set_cell_bg(cell, "0A1F44")
        p  = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r  = p.add_run(h)
        r.bold           = True
        r.font.size      = Pt(10)
        r.font.color.rgb = WHITE
    # data rows
    for ri, row_data in enumerate(rows):
        row = t.rows[ri + 1]
        bg  = "F7FAFD" if ri % 2 == 0 else "FFFFFF"
        for ci, cell_text in enumerate(row_data):
            cell = row.cells[ci]
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            # bold the first column entry
            r = p.add_run(cell_text)
            r.font.size = Pt(10)
            if ci == 0:
                r.bold           = True
                r.font.color.rgb = NAVY
            else:
                r.font.color.rgb = BLACK
    # column widths
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = w
    doc.add_paragraph()   # spacer


# ══════════════════════════════════════════════════════════════════════════════
#  COVER PAGE
# ══════════════════════════════════════════════════════════════════════════════
doc_title(doc, "DHRUVA CDM")
doc_subtitle(doc, "Autonomous Constellation Management & Conjunction Decision Support Platform")
doc_subtitle(doc, "Hackathon Presentation — Slide Content Document")
add_hr(doc)
note(doc, "Stack: Python · FastAPI · React · Next.js · WebSocket · physics.py")
note(doc, "Repository: github.com/jeevankumar-m/project-dhruva")
note(doc, "Demo Video: https://youtu.be/fs9eEaVNgIA")
note(doc, "Version 1.0.0  |  © 2026 Dhruva CDM Project")
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  SLIDE 1
# ══════════════════════════════════════════════════════════════════════════════
slide_label(doc, 1, "The Problem: Earth's Orbit Is Becoming a Minefield")
body(doc, "Suggested Slide Title (on PPT):", ["Suggested Slide Title (on PPT):"])
punchline(doc, '"The Kessler Syndrome Is Not a Theory — It\'s a Timeline"')

section_heading(doc, "Key Points")
bullet(doc, "Over 27,000 tracked debris objects currently orbit Earth — estimates suggest 500,000+ fragments above 1 cm, each carrying enough kinetic energy to destroy a satellite.", bold_lead="Over 27,000 tracked debris objects")
bullet(doc, "A single hypervelocity collision at LEO (7–8 km/s) generates thousands of new fragments, triggering a cascade — Kessler Syndrome — that can render entire orbital shells permanently unusable.", bold_lead="A single hypervelocity collision")
bullet(doc, "India's stakes are high: ISRO operates NavIC, Cartosat, RISAT, and emerging private constellations. Each is exposed every orbit.", bold_lead="India's stakes are high:")
bullet(doc, "Current global SSA tools — NASA CARA, ESA SST, LeoLabs — are siloed, expensive, and inaccessible to academic / startup operators.", bold_lead="Current global SSA tools")
bullet(doc, "No open-source, full-stack, physics-accurate CDM platform exists that a university lab, ISRO spinoff, or NewSpace company can deploy today.", bold_lead="No open-source, full-stack, physics-accurate CDM platform exists")

section_heading(doc, "Visual Suggestion")
note(doc, "A timeline graphic showing debris growth from 1957 → 2024, or the Iridium-Cosmos 2009 collision visualization. Numbers should dominate the slide visually.")
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  SLIDE 2
# ══════════════════════════════════════════════════════════════════════════════
slide_label(doc, 2, "Introducing Dhruva CDM")
punchline(doc, '"Dhruva — the North Star. The fixed point everything navigates by."')

section_heading(doc, "What Is It?")
body(doc,
     "Dhruva CDM is a full-stack Orbital Conjunction Data Management system that autonomously "
     "detects, assesses, and responds to satellite-debris close-approach events in real time. "
     "Built from first principles — not a wrapper around someone else's propagator. "
     "Every equation is implemented, validated, and running.")

section_heading(doc, "Three Coordinated Capabilities")
bullet(doc, "SENSE — Continuously propagate and track all objects in ECI (Earth-Centred Inertial) space.", bold_lead="SENSE")
bullet(doc, "ASSESS — Predict Time of Closest Approach (TCA) and classify risk across SAFE / WARNING / CRITICAL thresholds.", bold_lead="ASSESS")
bullet(doc, "ACT — Autonomously schedule and execute avoidance maneuvers with full operational constraint validation.", bold_lead="ACT")

section_heading(doc, "Deployment")
code_block(doc, "docker build -t dhruva-cdm . && docker run --rm -p 8000:8000 -p 3000:3000 dhruva-cdm")
body(doc, "One command. Running in under 60 seconds. Open source under MIT License.")

section_heading(doc, "Quick Links")
bullet(doc, "Demo Video: https://youtu.be/fs9eEaVNgIA")
bullet(doc, "Repository: github.com/jeevankumar-m/project-dhruva")
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  SLIDE 3
# ══════════════════════════════════════════════════════════════════════════════
slide_label(doc, 3, "Our USP — What Makes Dhruva Different")
punchline(doc, "Six Differentiators That No Existing Open Tool Combines")

add_table(doc,
    ["Dhruva USP", "Why It Matters"],
    [
        ["Physics-first propagation from scratch",
         "Two-body + J2 + RK4 — not SGP4 table lookups. J2 changes RAAN by ~7°/day; ignoring it means wrong conjunction geometry hours later."],
        ["Dual-mode conjunction pipeline",
         "Spatial Grid (O(1)/tick, UI-safe) + custom hand-built KD-Tree (sub-linear, stress-tested to 20,000 debris). No existing free tool does both."],
        ["Closed-loop autonomous avoidance",
         "CRITICAL events auto-generate evasion + paired recovery burns. No human in the loop required."],
        ["Full operational constraint model",
         "LOS blackout gating, thermal cooldown (600 s), command latency (10 s min), Tsiolkovsky fuel validation — all enforced before a burn is accepted."],
        ["WebSocket-streamed, sub-tick warning latency",
         "Risk state reaches the operator display within one simulation tick — not the next poll interval. Eliminates stale-state entirely."],
        ["End-of-life graveyard routing",
         "Satellites below 5% fuel are auto-flagged and routed to graveyard orbit — the IADC-compliant disposal lifecycle, fully automated."],
    ],
    col_widths=[Inches(2.4), Inches(4.0)]
)

punchline(doc, '"ESA\'s SST cost €100M+ to build. Dhruva CDM is open-source, dockerised, and running in under 60 seconds."')
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  SLIDE 4
# ══════════════════════════════════════════════════════════════════════════════
slide_label(doc, 4, "Physics Engine — The Scientific Core")
punchline(doc, '"Rigorous Astrodynamics — Every Number Has a Citation"')

section_heading(doc, "1. Two-Body Gravitational Acceleration")
code_block(doc, "a_grav = -(μ / |r|³) · r      where μ = 398,600.4418 km³/s²")
body(doc, "Dominant acceleration term for all Earth-orbiting objects. The baseline from which all perturbations are added.")

section_heading(doc, "2. J2 Oblateness Perturbation  (critically important — most student tools skip this)")
code_block(doc, "s = (3/2) · J2 · μ · R_E² / |r|⁵\nax = s·x·(5z²/|r|² - 1),  ay = s·y·(5z²/|r|² - 1),  az = s·z·(5z²/|r|² - 3)")
body(doc, "Earth's equatorial bulge causes RAAN to precess ~7°/day at ISS altitude. Without J2, conjunction predictions diverge from truth within hours. Dhruva includes the full zonal harmonic correction per Vallado (2013).")

section_heading(doc, "3. Fourth-Order Runge–Kutta Integration")
code_block(doc, "S_(n+1) = S_n + (Δt/6)·(k1 + 2k2 + 2k3 + k4)\nLocal truncation error O(Δt⁵)  |  Global error O(Δt⁴)  |  Fixed step dt = 10 s")
body(doc, "Validated against known ISS orbital period — accuracy < 0.1% error for eccentricities below 0.01.")

section_heading(doc, "4. Tsiolkovsky Propellant Model")
code_block(doc, "m_prop = m_current · (1 − exp(−|Δv| / Isp·g₀))    Isp = 300 s,  g₀ = 9.80665 m/s²")
body(doc, "Pre-validates every commanded maneuver against remaining fuel budget. Physically impossible burns are rejected before scheduling.")

section_heading(doc, "Reference Stack")
bullet(doc, "Vallado, Fundamentals of Astrodynamics and Applications, 4th ed., 2013")
bullet(doc, "Hoots & Roehrich, Spacetrack Report No. 3, 1980 (SGP4)")
bullet(doc, "NASA CARA, NPR 8715.6A, Debris Avoidance Maneuver Planning, 2017")
bullet(doc, "Alfano, Satellite Conjunction Monte Carlo Analysis, AAS 09-233, 2009")
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  SLIDE 5
# ══════════════════════════════════════════════════════════════════════════════
slide_label(doc, 5, "Conjunction Assessment Pipeline")
punchline(doc, '"From 20,000 Debris Objects to a Ranked Risk List — in One Tick"')

body(doc,
     "Naïve all-pairs assessment is O(N²). With 200 satellites × 20,000 debris = 4 million pairs per tick. "
     "Dhruva solves this with a two-stage, dual-mode pipeline.")

section_heading(doc, "Stage 1 — Candidate Pre-filter (two modes)")
bullet(doc, "Spatial Grid (live stream mode): Debris bucketed into 500 km³ cells. Each satellite checks only adjacent cells — bounded O(1) per-tick cost regardless of debris count. Protects UI responsiveness at 50+ Hz update rate.", bold_lead="Spatial Grid (live stream mode):")
bullet(doc, "Custom KD-Tree (batch / stress mode): Hand-built 3D KD-tree (zero scipy dependency). Radius query per satellite — sub-linear O(log N + k) scaling. Validated at 20,000 debris objects.", bold_lead="Custom KD-Tree (batch / stress mode):")

section_heading(doc, "Stage 2 — TCA Prediction")
bullet(doc, "Both objects propagated forward over a 2-hour look-ahead horizon at 30-second RK4 samples.", bold_lead="Both objects propagated forward")
bullet(doc, "Minimum separation recorded → Time of Closest Approach (TCA) and miss distance computed per pair.", bold_lead="Minimum separation recorded")

section_heading(doc, "Stage 3 — Risk Labelling")
add_table(doc,
    ["Risk Level", "Miss Distance", "System Response"],
    [
        ["CRITICAL", "< 1.0 km",  "Auto-evade burn triggered immediately + recovery burn queued"],
        ["WARNING",  "< 5.0 km",  "Operator alerted; optional manual maneuver"],
        ["SAFE",     "> 5.0 km",  "Continue monitoring; no action required"],
    ],
    col_widths=[Inches(1.3), Inches(1.5), Inches(3.6)]
)

body(doc, "Output: Up to 500 ranked conjunction warnings per snapshot, sorted by miss distance and TCA.")
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  SLIDE 6
# ══════════════════════════════════════════════════════════════════════════════
slide_label(doc, 6, "Autonomous Maneuver Planning & Validation")
punchline(doc, '"No Burn Flies Without Passing a 4-Gate Safety Gauntlet"')

section_heading(doc, "Validation Gate Sequence")
code_block(doc,
"""[Command Received]
      ↓
[1] Command Latency Check    → Burn epoch must be ≥ 10 s in the future (uplink time)
      ↓
[2] Thermal Cooldown Gate    → Minimum 600 s between consecutive burns (thruster protection)
      ↓
[3] Ground Station LOS       → Satellite must be visible to a ground station at burn epoch
                               (IIT Delhi Ground Node | ISTRAC Bengaluru | Goldstone Tracking)
      ↓
[4] Fuel Feasibility         → Tsiolkovsky check: |Δv| achievable with remaining propellant
      ↓
[SCHEDULED ✓  /  REJECTED ✗  with reason code]""")

section_heading(doc, "Autonomous CRITICAL Response")
bullet(doc, "Conjunction engine detects miss distance < 1 km → threshold crossed.", bold_lead="Conjunction engine detects")
bullet(doc, "Generates AUTO-EVA-* burn (velocity-direction impulse, auto-calculated magnitude).", bold_lead="Generates AUTO-EVA-*")
bullet(doc, "Immediately pairs it with AUTO-REC-* recovery burn to restore original orbital plane after conjunction window passes.", bold_lead="Immediately pairs it with AUTO-REC-*")
bullet(doc, "Full audit trail written to burn logs — every burn timestamped, satellite identified, ΔV recorded, status marked.", bold_lead="Full audit trail")

section_heading(doc, "Live Demo Result")
bullet(doc, "15 burns executed (AUTO-EVA, AUTO-REC, operator EVASION, operator RECOVERY).")
bullet(doc, "1 burn correctly rejected — EVASION-01 for SAT-Alpha-01 at 15.17 m/s — LOS constraint violated.")
bullet(doc, "The rejection proves the system is safe-by-default, not just happy-path correct.")
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  SLIDE 7
# ══════════════════════════════════════════════════════════════════════════════
slide_label(doc, 7, "System Architecture")
punchline(doc, '"Service-Oriented, Streaming, Production-Grade — From Day One"')

code_block(doc,
"""┌──────────────────────────────────────────────────────────────┐
│           OPERATOR FRONTEND  (React / Next.js)               │
│  2D Map  │  3D Cesium  │  Bullseye  │  Analytics  │  CDM     │
└───────────────────────┬──────────────────────────────────────┘
                        │  WebSocket /orbit  (live stream)
                        │  REST /api/*       (commands)
┌───────────────────────▼──────────────────────────────────────┐
│           FASTAPI BACKEND  (server.py)                       │
│  api/routes.py  │  engine/sim_state.py   │  physics.py       │
│                 │  engine/conjunction.py  (Grid / KD-Tree)   │
└──────────────────────────────────────────────────────────────┘
      INPUT:  TLE Data  |  Cartesian Telemetry  |  CDM CSV""")

section_heading(doc, "Key Architectural Decisions")
bullet(doc, "WebSocket over REST polling: eliminates stale-state and synchronisation artifacts. Warning state reaches the UI within a single simulation tick (sub-second).", bold_lead="WebSocket over REST polling:")
bullet(doc, "Dual execution modes: Live stream (spatial grid, UI-optimised) vs. Step/Batch (KD-tree, throughput-optimised). One backend, two operating profiles.", bold_lead="Dual execution modes:")
bullet(doc, "TLE → SGP4 → ECI pipeline: ingests real NORAD/CelesTrak catalog data; backend handles frame transformation automatically (TEME → GCRS).", bold_lead="TLE → SGP4 → ECI pipeline:")
bullet(doc, "Time-warp 1×–100×: compress orbital mechanics for rapid scenario rehearsal and stress testing.", bold_lead="Time-warp 1×–100×:")
bullet(doc, "Single-container Docker deploy: backend (port 8000) + frontend (port 3000) in one image. Zero-dependency setup for evaluators.", bold_lead="Single-container Docker deploy:")

section_heading(doc, "Module Responsibilities")
add_table(doc,
    ["Module", "Responsibility"],
    [
        ["server.py",           "FastAPI bootstrap, CORS, SimulationState instantiation, WebSocket /orbit endpoint"],
        ["api/routes.py",       "All REST endpoints: telemetry ingest, TLE ingest, maneuver scheduling, step, snapshot, seed, time-warp, CDM CSV loader"],
        ["engine/sim_state.py", "Operational core: sim time, state vectors, maneuver queues, conjunction warnings, blackout tracking, graveyard logic, objective metrics"],
        ["engine/conjunction.py","Conjunction pipeline: grid + KD-tree pre-filter, TCA evaluation, SAFE/WARNING/CRITICAL classification"],
        ["physics.py",          "Astrodynamics: two-body + J2, RK4 integration, state propagation, Tsiolkovsky propellant model, orbital period"],
        ["dhruva-frontend/",    "React SPA: 2D map, 3D Cesium globe, Bullseye radar, analytics dashboard, maneuver workflow, burn logs"],
    ],
    col_widths=[Inches(1.8), Inches(4.6)]
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  SLIDE 8
# ══════════════════════════════════════════════════════════════════════════════
slide_label(doc, 8, "Operator Dashboard — Orbital Insight")
punchline(doc, '"Six Views. One Ground Truth. Zero Polling Lag."')

add_table(doc,
    ["Dashboard View", "What It Delivers"],
    [
        ["2D Mercator Ground Track",
         "Real-time satellite positions (cyan), debris cloud (red), 90-min historical trail + 90-min predicted trajectory (dashed), solar terminator line, ground station overlays (IIT Delhi Node, ISTRAC, Goldstone)"],
        ["3D Cesium Globe",
         "Full-constellation 3D render — polar view, horizon view, flat-Earth spread — with live blackout flagging and per-satellite telemetry panel (fuel mass, altitude, longitude, drift-from-slot)"],
        ["Conjunction Bullseye Radar",
         "Polar radar plot per satellite: debris approach vector encoded as angle, TCA as radial distance, colour-coded risk rings (1 km / 2 km / 5 km). Expandable modal with full debris threat list."],
        ["Fleet Fuel & ΔV Analytics",
         "Per-satellite propellant gauges for all 50 satellites, cumulative ΔV consumed vs. collisions avoided (rolling chart), uptime score, avoidance efficiency metric"],
        ["Maneuver Timeline (Gantt)",
         "Scheduled burns, executed burns (green), cooldown windows (teal), rejected commands (red) per satellite across simulation timeline — full scheduling situational awareness"],
        ["CDM Alerts + Burn Logs",
         "Active conjunction warnings sorted by miss distance and TCA; complete audit trail of all AUTO-EVA-*, AUTO-REC-*, and operator-initiated burns with status codes"],
    ],
    col_widths=[Inches(2.0), Inches(4.4)]
)

section_heading(doc, "Graveyard Lifecycle (automated)")
code_block(doc, "NOMINAL  →  GRAVEYARD CANDIDATE  →  GRAVEYARD MANEUVER SCHEDULED  →  GRAVEYARD")
note(doc, "Trigger: fuel drops below 5% of initial propellant mass. Disposal route: above GEO (>36,000 km) for high-altitude missions; below ISS altitude (<400 km, re-entry decay) for LEO missions.")
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  SLIDE 9
# ══════════════════════════════════════════════════════════════════════════════
slide_label(doc, 9, "Scalability & Stress Testing")
punchline(doc, '"200 Satellites. 20,000 Debris. Sub-Second Updates. Verified."')

section_heading(doc, "Validated Operational Limits")
add_table(doc,
    ["Parameter", "Limit / Value", "Notes"],
    [
        ["Satellites",                  "Up to 200",           "Tested at 50-satellite baseline with full conjunction pipeline"],
        ["Debris objects",              "Up to 20,000",        "CDM CSV loader pipeline validated end-to-end"],
        ["Conjunction warnings/snapshot","Up to 500",          "Sorted by miss distance and TCA; worst-case captured first"],
        ["Time-warp multipliers",       "1×, 2×, 10×, 25×, 100×", "Scenario compression for rapid rehearsal"],
        ["TCA look-ahead horizon",      "7,200 s (2 hours)",   "At 30-second RK4 sample interval per object pair"],
        ["WebSocket update rate",       "Sub-second per tick", "Spatial grid: bounded per-tick cost regardless of debris count"],
        ["Max burn magnitude",          "15.0 m/s per burn",   "Enforced cap; prevents runaway fuel consumption"],
    ],
    col_widths=[Inches(2.2), Inches(1.6), Inches(2.6)]
)

section_heading(doc, "Why the Custom KD-Tree Matters at Scale")
bullet(doc, "Grid scan at 20,000 debris = large cell iteration count per satellite. Performance degrades linearly.", bold_lead="Grid scan at 20,000 debris")
bullet(doc, "KD-tree radius query scales as O(log N + k) — dramatically faster for uniformly distributed populations.", bold_lead="KD-tree radius query scales as O(log N + k)")
bullet(doc, "Custom implementation (zero scipy, zero external dependency) — portable to edge compute or onboard systems.", bold_lead="Custom implementation (zero scipy)")
bullet(doc, "Rebuilt each conjunction assessment call — handles dynamic debris states correctly (no stale index).", bold_lead="Rebuilt each conjunction assessment call")

section_heading(doc, "CDM CSV Stress Pipeline")
body(doc,
     "Loads RTN-relative state fields from CCSDS-style CDM CSV files, anchors them to in-simulation satellite primaries, "
     "converts to ECI-like Cartesian state vectors, and injects seamlessly into the conjunction pipeline. "
     "10,000 objects loaded, assessed, and warnings generated in a single API call.")
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  SLIDE 10
# ══════════════════════════════════════════════════════════════════════════════
slide_label(doc, 10, "Real-World Applicability")
punchline(doc, '"Three Deployment Scenarios Where Dhruva CDM Fills a Real Gap"')

section_heading(doc, "Scenario 1 — ISRO / NewSpace Operators")
body(doc,
     "ISRO's PSLV-C58 placed 10 satellites in a 650 km orbit — the exact LEO regime Dhruva CDM targets. "
     "A Flight Dynamics Officer using Dhruva gets real-time CDM alerts, autonomous avoidance, and a full burn-log audit "
     "without a $10M+ commercial SSA subscription. TLE ingestion from CelesTrak requires a single API call.")

section_heading(doc, "Scenario 2 — University Constellation Labs")
body(doc,
     "IIT Delhi's Student Satellite Programme, IIST, BITS — student teams operating CubeSats have no CDM tooling available to them. "
     "Dhruva CDM deploys in a Docker container on a laptop, runs on real TLE data, and gives students the same "
     "situational awareness a Flight Dynamics Officer has — for free.")

section_heading(doc, "Scenario 3 — Inter-Agency Space Traffic Management")
body(doc,
     "The CCSDS CDM message standard is the basis for NASA–ESA–JAXA data exchange. "
     "Dhruva CDM's architecture is explicitly designed for extension to CCSDS format (roadmap), enabling it to become "
     "a node in a federated, multi-agency STM network.")

section_heading(doc, "Alignment with National Priorities")
bullet(doc, "India Space Policy 2023 mandates national SSA capability development — Dhruva CDM is a direct contribution.", bold_lead="India Space Policy 2023")
bullet(doc, "ISRO's SSA Control Centre (Bengaluru) — Dhruva CDM's architecture maps directly to its operational mission.", bold_lead="ISRO's SSA Control Centre (Bengaluru)")
bullet(doc, "IN-SPACe framework encourages open-source tooling for commercial space operators — Dhruva CDM is MIT-licensed.", bold_lead="IN-SPACe framework")
bullet(doc, "IIT Delhi Ground Station is pre-configured as a named ground station node in Dhruva CDM.", bold_lead="IIT Delhi Ground Station")
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  SLIDE 11
# ══════════════════════════════════════════════════════════════════════════════
slide_label(doc, 11, "Live Demo Highlights & Key Results")
punchline(doc, '"The Numbers from a Live Simulation Run"')

section_heading(doc, "Simulation Parameters")
bullet(doc, "50 active satellites  |  20 seeded debris  +  10,000 CDM CSV debris loaded (20,020 total tracked objects)")
bullet(doc, "Time-warp: variable (1× – 100×)  |  Look-ahead horizon: 2 hours  |  RK4 sample: 30 s")

section_heading(doc, "Mission Metrics (from live run)")
add_table(doc,
    ["Metric", "Result"],
    [
        ["Total ΔV consumed (fleet)",        "124.11 m/s"],
        ["Fleet uptime",                     "92.74%"],
        ["Uptime score",                     "80.42"],
        ["Outage satellite-seconds",         "18,206 s"],
        ["Total burns logged",               "16 (15 executed / 1 correctly rejected)"],
        ["Active CDM warnings (live sample)","3 SAFE — 7.636 km (TCA 1080 s), 14.582 km (TCA 1800 s), 36.157 km (TCA 60 s)"],
        ["Graveyard satellites",             "0 — propellant budgeting maintained entire 50-satellite fleet above EOL threshold"],
        ["Max simultaneous blackouts",       "12 satellites (tracked and surfaced to operator simultaneously)"],
    ],
    col_widths=[Inches(3.0), Inches(3.4)]
)

section_heading(doc, "The Genesis Story (Origin Proof-of-Concept)")
bullet(doc, "Stage 1: 2D simulation — predicted collision point identified on look-ahead horizon. Satellite (green) and debris (red) on separate tracks.", bold_lead="Stage 1:")
bullet(doc, "Stage 2: Orbits converged, predicted collision point migrated toward satellite position — threshold crossed — evasion burn scheduled.", bold_lead="Stage 2:")
bullet(doc, "Stage 3: Post-burn trajectory confirmed diverged from debris track — closed-loop autonomy validated in the very first prototype.", bold_lead="Stage 3:")

section_heading(doc, "The Rejected Burn — Why It Matters")
body(doc,
     "EVASION-01 for SAT-Alpha-01 (15.17 m/s) was rejected by the LOS/cooldown validation gate. "
     "This is not a failure — it is proof that the system is safe-by-default. A burn that would violate "
     "an operational constraint is rejected with a reason code, not silently ignored. That is production-grade behaviour.")
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  SLIDE 12
# ══════════════════════════════════════════════════════════════════════════════
slide_label(doc, 12, "Roadmap & Vision")
punchline(doc, '"Dhruva CDM v2.0 — The Path to Operational Readiness"')

section_heading(doc, "Immediate Extensions (3–6 months)")
bullet(doc, "Live NORAD TLE ingestion (CelesTrak / Space-Track.org) — automated SGP4 seeding from real orbital catalog; real objects, real risk assessments.", bold_lead="Live NORAD TLE ingestion")
bullet(doc, "Probabilistic collision avoidance (Pc) — replace threshold-based labelling with full covariance Foster/Alfano 2D projection method, matching the NASA CARA Pc > 10⁻⁴ operational threshold standard.", bold_lead="Probabilistic collision avoidance (Pc)")

section_heading(doc, "Medium-Term Research Track (6–18 months)")
bullet(doc, "Higher-fidelity propagation — J3–J6 zonal harmonics, NRLMSISE-00 atmospheric drag, solar radiation pressure → extends accurate propagation from hours to days; critical for VLEO objects.", bold_lead="Higher-fidelity propagation")
bullet(doc, "Reinforcement Learning maneuver planner — policy trained on simulated conjunction scenarios, minimising fleet-wide ΔV while maximising uptime and respecting thermal cooldown constraints (Gao et al., arXiv:2009.04050).", bold_lead="Reinforcement Learning maneuver planner")

section_heading(doc, "Long-Term Vision")
bullet(doc, "CCSDS CDM standard API — interoperability with NASA, ESA, JAXA, ISRO data exchange frameworks → Dhruva as a node in federated Space Traffic Management.", bold_lead="CCSDS CDM standard API")
bullet(doc, "Onboard autonomy migration — evasion planner packaged as a radiation-tolerant flight software module with deterministic WCET guarantees → fully autonomous onboard avoidance for satellites in prolonged blackout.", bold_lead="Onboard autonomy migration")

add_hr(doc)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(14)
r = p.add_run("The star Dhruva does not move. Everything else navigates around it.")
r.bold = True; r.italic = True
r.font.size = Pt(13); r.font.color.rgb = NAVY

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run("Our goal: build the fixed reference point for India's space safety infrastructure.")
r2.italic = True; r2.font.size = Pt(11); r2.font.color.rgb = BLUE
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  DELIVERY NOTES PAGE
# ══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
r = p.add_run("PRESENTER DELIVERY NOTES")
r.bold = True; r.font.size = Pt(14); r.font.color.rgb = NAVY
add_hr(doc)

notes_list = [
    ("Slide 1 — Open Hard",
     "Quote the fragment count (500,000+) and let it land before moving on. Do not rush past it. "
     "Let the judges absorb the scale of the problem before you offer the solution."),
    ("Slide 3 — Say the Comparison Line Out Loud",
     '"ESA\'s SST cost over a hundred million euros. This runs in 60 seconds on a laptop." '
     "Pause after saying it. It will land."),
    ("Slide 4 — On Physics Depth",
     "You don't need to derive the math. Just say: "
     '"We implemented this from scratch — not a library call. Every equation is in physics.py and cites Vallado 2013." '
     "That alone impresses an ISRO judge."),
    ("Slide 5 — The O(N²) Line",
     "Say the number: 4 million pairs per tick. Then explain how the pipeline reduces it. "
     "Concrete numbers always beat abstract complexity notation for a general audience."),
    ("Slide 6 — Walk the Gauntlet Like a Checklist",
     "Read through the 4 gates slowly — Latency, Cooldown, LOS, Fuel. "
     "Let it sound procedurally rigorous. Judges from ISRO / DRDO will recognise this as flight-software thinking."),
    ("Slide 11 — Lead with the Rejected Burn",
     "Open Slide 11 with: 'The most important result isn't a burn that executed — it's the one we correctly rejected.' "
     "This flips the frame from feature showcase to safety assurance."),
    ("Slide 12 — Close with the Tagline",
     "The Sanskrit meaning of Dhruva landing at the very end of the presentation is a strong, "
     "memorable close. Deliver it slowly. Don't rush the Thank You slide after it."),
]

for title, content in notes_list:
    section_heading(doc, title)
    body(doc, content)

# ── Save ──────────────────────────────────────────────────────────────────────
out_path = r"C:\Users\Jeevan Kumar M\Desktop\Project-dhruva\Dhruva_CDM_PPT_Content.docx"
doc.save(out_path)
print("Saved: " + out_path)
